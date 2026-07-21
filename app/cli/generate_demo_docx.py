from pathlib import Path

from docx import Document


def generate(path: str | Path = "demo/demo_protocol.docx") -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Протокол демонстрационного совещания", 0)
    doc.add_paragraph("Номер: DEMO-001")
    doc.add_paragraph("Дата: 21.07.2026")
    for title in ["1. Запуск проекта", "2. Контроль качества"]:
        doc.add_heading(title, level=1)
        doc.add_paragraph("Обсудили автономный сценарий без внешних API.")
        doc.add_paragraph("1. Иванов Иван Иванович подготовить план работ до 31.07.2026.")
        doc.add_paragraph("2. Петров Петр Петрович оформить критерии приемки в течение 5 дней.")
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Поручение"
    hdr[2].text = "Исполнитель"
    hdr[3].text = "Срок"
    rows = [
        ("3", "Согласовать демонстрацию", "Сидоров С.С.; Иванов И.И.", "05.08.2026"),
        ("4", "Проверить неоднозначное ФИО", "Алексей", ""),
    ]
    for r in rows:
        cells = table.add_row().cells
        for i, value in enumerate(r):
            cells[i].text = value
    doc.save(path)
    print(path)
    return path


if __name__ == "__main__":
    generate()
