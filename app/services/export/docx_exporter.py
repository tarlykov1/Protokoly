from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from sqlalchemy import select
from sqlalchemy.orm import Session, selectinload

from app.db.models.domain import Protocol, ProtocolSection, ProtocolTask, ProtocolTaskAssignment


class ProtocolDocxExporter:
    """Export protocols either as a round-trip MEMO or as a reader-friendly document."""

    def __init__(self, db: Session):
        self.db = db

    def export(self, protocol_id: int, mode: str = "memo") -> bytes:
        if mode not in {"memo", "print"}:
            raise ValueError("Неизвестный режим экспорта")
        protocol = self.db.scalar(
            select(Protocol)
            .where(Protocol.id == protocol_id)
            .options(
                selectinload(Protocol.project),
                selectinload(Protocol.tasks)
                .selectinload(ProtocolTask.assignments)
                .selectinload(ProtocolTaskAssignment.employee),
            )
        )
        if protocol is None:
            raise ValueError("Протокол не найден")
        sections = self.db.scalars(
            select(ProtocolSection)
            .where(ProtocolSection.protocol_id == protocol_id)
            .order_by(ProtocolSection.sort_order, ProtocolSection.id)
        ).all()
        tasks = sorted(protocol.tasks, key=lambda task: (task.position, task.id))
        document = Document()
        document.styles["Normal"].font.name = "Arial"
        document.styles["Normal"].font.size = Pt(10)
        if mode == "memo":
            self._build_memo(document, protocol, sections, tasks)
        else:
            self._build_print(document, protocol, sections, tasks)
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    @staticmethod
    def _assignees(task: ProtocolTask) -> str:
        return ", ".join(
            assignment.assignee_name or "—"
            for assignment in sorted(task.assignments, key=lambda item: item.sort_order)
        ) or "—"

    def _build_memo(self, document, protocol, sections, tasks) -> None:
        """Use paragraphs only: this is the canonical MemoProtocolParser contract."""
        number = (protocol.number or str(protocol.id)).replace("M-", "М – ").replace("М-", "М – ")
        document.add_paragraph(number)
        heading = document.add_paragraph("ИТОГИ")
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].bold = True
        title = document.add_paragraph(protocol.title)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if protocol.meeting_date:
            document.add_paragraph(
                f"г. Санкт-Петербург «{protocol.meeting_date:%d}» "
                f"{self._month(protocol.meeting_date.month)} {protocol.meeting_date:%Y} года"
            )
        document.add_paragraph("ОТМЕТИЛИ:").runs[0].bold = True
        document.add_paragraph(protocol.description or "—")
        document.add_paragraph("РЕШИЛИ:").runs[0].bold = True
        task_groups = [(section.title, [t for t in tasks if t.section_id == section.id]) for section in sections]
        unsectioned = [t for t in tasks if not t.section_id]
        if unsectioned:
            task_groups.append(("Без раздела", unsectioned))
        task_index = 0
        for section_title, section_tasks in task_groups:
            document.add_paragraph(f"#{section_title}").runs[0].bold = True
            for task in section_tasks:
                task_index += 1
                # Numbering follows persisted order. The original task number remains editable in UI,
                # but contiguous MEMO numbering makes repeated imports deterministic.
                document.add_paragraph(f"{task_index}. {task.title}")
                document.add_paragraph("Исполнители:").runs[0].bold = True
                document.add_paragraph(self._assignees(task))
                document.add_paragraph("Срок:").runs[0].bold = True
                document.add_paragraph(task.deadline.strftime("%d.%m.%Y") if task.deadline else "Без срока")

    def _build_print(self, document, protocol, sections, tasks) -> None:
        heading = document.add_paragraph(protocol.title)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].bold = True
        document.add_paragraph(f"Протокол № {protocol.number or '—'}")
        for section in sections:
            document.add_heading(section.title, level=2)
            for task in [t for t in tasks if t.section_id == section.id]:
                paragraph = document.add_paragraph(style="List Number")
                paragraph.add_run(task.title).bold = True
                if task.description and task.description != task.title:
                    document.add_paragraph(task.description)
                document.add_paragraph(f"Исполнители: {self._assignees(task)}")
                document.add_paragraph(
                    f"Срок: {task.deadline:%d.%m.%Y}" if task.deadline else "Срок: —"
                )

    @staticmethod
    def _month(month: int) -> str:
        return ("января", "февраля", "марта", "апреля", "мая", "июня", "июля", "августа", "сентября", "октября", "ноября", "декабря")[month - 1]
