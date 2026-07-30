from __future__ import annotations

import re

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from app.parsers.protocol import ParsedDocument, ParsedElement, ParsedParagraph, SourceLocation


def _clean(text: str) -> str:
    text = text.replace("\xa0", " ").replace("\u202f", " ")
    return re.sub(r"[ \t]+", " ", text).strip()


def _paragraph_text(paragraph: Paragraph) -> str:
    return _clean("".join(run.text for run in paragraph.runs) or paragraph.text)


def _cell_text(cell: _Cell) -> str:
    return "\n".join(filter(None, (_paragraph_text(p) for p in cell.paragraphs)))


def iter_docx_elements(doc: DocxDocument):
    p_idx = 0
    t_idx = 0
    order = 0
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            text = _paragraph_text(p)
            if text:
                yield ParsedElement(
                    order,
                    text,
                    "paragraph",
                    SourceLocation(paragraph_index=p_idx),
                    p.style.name if p.style else None,
                )
                order += 1
            p_idx += 1
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for r_idx, row in enumerate(table.rows):
                # Keep empty cells: their positions are meaningful in forms where a label and
                # its value live in neighbouring cells.  A horizontally merged cell is exposed
                # by python-docx more than once, so only read its underlying XML node once.
                seen_cells: set[int] = set()
                values: list[str] = []
                for cell in row.cells:
                    cell_id = id(cell._tc)
                    values.append("" if cell_id in seen_cells else _cell_text(cell))
                    seen_cells.add(cell_id)
                cells = tuple(values)
                text = _clean(" | ".join(c for c in cells if c))
                if text:
                    yield ParsedElement(
                        order,
                        text,
                        "table_row",
                        SourceLocation(table_index=t_idx, row_index=r_idx),
                        None,
                        cells,
                    )
                    order += 1
            t_idx += 1


def parse_docx(path: str) -> ParsedDocument:
    doc = Document(path)
    elements = list(iter_docx_elements(doc))
    paragraphs = [
        ParsedParagraph(e.location.paragraph_index or 0, e.text)
        for e in elements
        if e.kind == "paragraph"
    ]
    return ParsedDocument(filename=path, paragraphs=paragraphs, elements=elements)
