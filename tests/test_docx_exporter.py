from datetime import date
from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import delete, select

from app.db.base import Base
from app.db.models.domain import (
    Employee,
    Project,
    Protocol,
    ProtocolSection,
    ProtocolTask,
    ProtocolTaskAssignment,
)
from app.db.session import SessionLocal, engine
from app.main import app
from app.services.export import ProtocolDocxExporter


def make_export_protocol():
    Base.metadata.create_all(engine)
    with SessionLocal() as db:
        project = Project(name="Export", code="EXPORT-DOCX")
        first_employee = Employee(full_name="Петров Петр Петрович", is_active=True)
        second_employee = Employee(full_name="Сидорова Анна Игоревна", is_active=True)
        db.add_all([project, first_employee, second_employee])
        db.flush()
        protocol = Protocol(
            project_id=project.id,
            title="МЕМО по проекту",
            number="M-026/26",
            meeting_date=date(2026, 8, 5),
            initiator="Офис проекта",
            responsible="PMO",
        )
        db.add(protocol)
        db.flush()
        section = ProtocolSection(protocol_id=protocol.id, title="Решения", sort_order=1)
        db.add(section)
        db.flush()
        first = ProtocolTask(
            protocol_id=protocol.id,
            section_id=section.id,
            number="2",
            position=20,
            title="Второе поручение",
            deadline=date(2026, 9, 2),
        )
        second = ProtocolTask(
            protocol_id=protocol.id,
            section_id=section.id,
            number="1",
            position=10,
            title="Первое поручение",
            deadline=date(2026, 9, 1),
        )
        db.add_all([first, second])
        db.flush()
        db.add_all([
            ProtocolTaskAssignment(protocol_task_id=first.id, employee_id=second_employee.id),
            ProtocolTaskAssignment(protocol_task_id=second.id, employee_id=first_employee.id),
        ])
        db.commit()
        return protocol.id


def teardown_export_protocol():
    with SessionLocal() as db:
        project = db.scalar(select(Project).where(Project.code == "EXPORT-DOCX"))
        if project:
            protocol_ids = select(Protocol.id).where(Protocol.project_id == project.id)
            task_ids = select(ProtocolTask.id).where(ProtocolTask.protocol_id.in_(protocol_ids))
            db.execute(delete(ProtocolTaskAssignment).where(ProtocolTaskAssignment.protocol_task_id.in_(task_ids)))
            db.execute(delete(ProtocolTask).where(ProtocolTask.protocol_id.in_(protocol_ids)))
            db.execute(delete(ProtocolSection).where(ProtocolSection.protocol_id.in_(protocol_ids)))
            db.execute(delete(Protocol).where(Protocol.project_id == project.id))
            db.execute(delete(Employee).where(Employee.full_name.in_(["Петров Петр Петрович", "Сидорова Анна Игоревна"])))
            db.execute(delete(Project).where(Project.id == project.id))
            db.commit()


def document_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + table_cells)


def test_docx_exporter_creates_file_with_tasks_assignees_deadlines_and_order():
    teardown_export_protocol()
    protocol_id = make_export_protocol()
    try:
        with SessionLocal() as db:
            content = ProtocolDocxExporter(db).export(protocol_id)

        assert content.startswith(b"PK")
        text = document_text(content)
        assert "МЕМО по проекту" in text
        assert "Первое поручение" in text
        assert "Второе поручение" in text
        assert "Петров Петр Петрович" in text
        assert "02.09.2026" in text
        assert text.index("Первое поручение") < text.index("Второе поручение")
    finally:
        teardown_export_protocol()


def test_docx_export_endpoint_returns_attachment_file():
    teardown_export_protocol()
    protocol_id = make_export_protocol()
    try:
        response = TestClient(app).get(f"/protocols/{protocol_id}/export/docx")

        assert response.status_code == 200
        assert response.headers["content-type"].startswith("application/vnd.openxmlformats-officedocument")
        assert 'filename="protocol_M-026_26.docx"' in response.headers["content-disposition"]
        assert "Первое поручение" in document_text(response.content)
    finally:
        teardown_export_protocol()


def test_memo_export_round_trip_preserves_structure(tmp_path):
    """The default download is a parser contract, not merely a visual report."""
    from app.parsers.docx import parse_docx
    from app.parsers.protocol import MemoProtocolParser

    Base.metadata.create_all(engine)
    teardown_export_protocol()
    protocol_id = make_export_protocol()
    try:
        with SessionLocal() as db:
            content = ProtocolDocxExporter(db).export(protocol_id, mode="memo")
        path = tmp_path / "memo-roundtrip.docx"
        path.write_bytes(content)
        document = Document(BytesIO(content))
        assert not document.tables

        result = MemoProtocolParser().parse(parse_docx(str(path)))

        assert result.errors == []
        assert [section.title for section in result.sections] == ["Решения"]
        assert [task.title for task in result.tasks] == ["Первое поручение", "Второе поручение"]
        assert [task.assignee_raw for task in result.tasks] == [
            "Петров Петр Петрович", "Сидорова Анна Игоревна"
        ]
        assert [task.deadline for task in result.tasks] == ["2026-09-01", "2026-09-02"]
    finally:
        teardown_export_protocol()


def test_print_and_memo_export_modes_are_available():
    Base.metadata.create_all(engine)
    teardown_export_protocol()
    protocol_id = make_export_protocol()
    try:
        client = TestClient(app)
        memo = client.get(f"/protocols/{protocol_id}/export/docx?mode=memo")
        printable = client.get(f"/protocols/{protocol_id}/export/docx?mode=print")
        assert memo.status_code == printable.status_code == 200
        assert "РЕШИЛИ:" in document_text(memo.content)
        assert "Протокол №" in document_text(printable.content)
    finally:
        teardown_export_protocol()
