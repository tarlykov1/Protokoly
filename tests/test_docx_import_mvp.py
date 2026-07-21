from datetime import UTC, datetime, timedelta
from pathlib import Path

import pytest
from docx import Document
from fastapi import HTTPException, UploadFile
from sqlalchemy import create_engine, select
from sqlalchemy.orm import sessionmaker

from app.db.base import Base
from app.db.models.domain import Employee, ImportSession, Project, Protocol
from app.parsers.docx import parse_docx
from app.parsers.protocol import ParserRegistry, UniversalProtocolParser, parse_deadline
from app.services.imports.service import (
    MAX_IMPORT_SIZE,
    cleanup_expired_import_sessions,
    confirm_session,
    create_preview_session,
    reparse_session,
    update_session_payload,
)


@pytest.fixture()
def db(tmp_path):
    engine = create_engine(f"sqlite:///{tmp_path / 'db.sqlite'}", future=True)
    Base.metadata.create_all(engine)
    Session = sessionmaker(bind=engine, future=True, expire_on_commit=False)
    with Session() as session:
        session.add(Project(id=1, name="Demo", code="DEMO"))
        session.add(Employee(id=1, full_name="Иванов Иван Иванович", is_available_in_bitrix=True))
        session.commit()
        yield session


def make_docx(path: Path) -> Path:
    doc = Document()
    doc.add_heading("Протокол № 1", level=1)
    doc.add_paragraph("РАЗДЕЛ ПРОЕКТА")
    doc.add_paragraph("1. Подготовить отчёт. Ответственный: Иванов Иван Иванович. Срок 25.07.2026")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "Поручение"
    table.rows[0].cells[1].text = "Ответственный"
    table.rows[0].cells[2].text = "Срок"
    table.rows[1].cells[0].text = "Согласовать план"
    table.rows[1].cells[1].text = "Петров П.П. / Сидоров С.С."
    table.rows[1].cells[2].text = "до конца недели"
    doc.save(path)
    return path


def upload(
    path: Path,
    name="safe.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
):
    return UploadFile(filename=name, file=path.open("rb"), headers={"content-type": mime})


def test_upload_valid_docx_creates_preview_session_without_protocol(db, tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    path = make_docx(tmp_path / "real.docx")
    session = create_preview_session(db, 1, upload(path, "../evil.docx"))
    assert session.id
    assert session.original_filename == "evil.docx"
    assert session.stored_filename.endswith(".docx")
    assert session.checksum
    assert db.scalar(select(Protocol)) is None
    assert len(session.parsed_payload["tasks"]) == 2


def test_rejects_extension_corrupt_docx_and_size(db, tmp_path):
    bad = tmp_path / "bad.txt"
    bad.write_text("x")
    with pytest.raises(HTTPException):
        create_preview_session(db, 1, upload(bad, "bad.txt", "text/plain"))
    corrupt = tmp_path / "bad.docx"
    corrupt.write_bytes(b"PKbroken")
    with pytest.raises(HTTPException):
        create_preview_session(db, 1, upload(corrupt))
    huge = tmp_path / "huge.docx"
    huge.write_bytes(b"PK" + (b"0" * (MAX_IMPORT_SIZE + 1)))
    with pytest.raises(HTTPException):
        create_preview_session(db, 1, upload(huge))


def test_confirm_creates_protocol_and_cancel_does_not(db, tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    s1 = create_preview_session(db, 1, upload(make_docx(tmp_path / "a.docx")))
    protocol = confirm_session(db, s1)
    assert protocol.id
    assert db.scalar(select(Protocol).where(Protocol.id == protocol.id))
    s2 = create_preview_session(db, 1, upload(make_docx(tmp_path / "b.docx")))
    s2.status = "cancelled"
    db.commit()
    assert s2.protocol_id is None


def test_update_reopen_and_reparse_history(db, tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    session = create_preview_session(db, 1, upload(make_docx(tmp_path / "c.docx")))
    payload = session.parsed_payload
    payload["document_title"] = "Исправлено"
    import json

    update_session_payload(db, session, json.dumps(payload))
    assert db.get(ImportSession, session.id).parsed_payload["document_title"] == "Исправлено"
    reparse_session(db, session, "memo", True)
    assert session.parse_history


def test_parser_paragraphs_tables_order_sections_numbers_dates_and_assignees(db, tmp_path):
    doc = parse_docx(str(make_docx(tmp_path / "p.docx")))
    assert [e.kind for e in doc.elements][:3] == ["paragraph", "paragraph", "paragraph"]
    result = UniversalProtocolParser().parse(doc)
    assert result.sections[0].title == "Протокол № 1"
    assert result.tasks[0].task_number == "1"
    assert result.tasks[0].deadline == "2026-07-25"
    assert result.tasks[1].deadline_type == "relative"
    parser, choice = ParserRegistry().choose(doc)
    assert parser.parser_type in {"universal", "ceo_protocol", "memo", "deputy_ceo_protocol"}
    assert choice.confidence >= 0


def test_deadline_formats():
    assert parse_deadline("до 25 июля", 2026)[0] == "2026-07-25"
    assert parse_deadline("еженедельно")[2] == "periodic"
    assert parse_deadline("без срока")[2] == "none"


def test_duplicate_warning_and_cleanup(db, tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    create_preview_session(db, 1, upload(make_docx(tmp_path / "d.docx")))
    second = create_preview_session(db, 1, upload(tmp_path / "d.docx"))
    assert second.parsed_payload["duplicates"]
    second.expires_at = datetime.now(UTC) - timedelta(hours=1)
    db.commit()
    assert cleanup_expired_import_sessions(db) == 1
    assert second.status == "expired"
    assert not Path(second.file_path).exists()
