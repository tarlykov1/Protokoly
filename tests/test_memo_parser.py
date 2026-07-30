import base64
from pathlib import Path

import pytest
from docx import Document
from fastapi import HTTPException
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from app.db.base import Base
from app.db.models.domain import Employee, ImportSession, Project
from app.parsers.docx import parse_docx
from app.parsers.protocol import (
    MemoProtocolParser,
    ParsedDocument,
    ParsedElement,
    ParserRegistry,
    SourceLocation,
    UniversalProtocolParser,
)
from app.services.imports.service import confirm_session, create_preview_session
from tests.test_docx_import_mvp import upload

FIXTURE_TEXT = Path("tests/fixtures/memo_m026_26.txt")


@pytest.fixture()
def db(tmp_path, monkeypatch):
    monkeypatch.chdir(Path(__file__).resolve().parents[1])
    engine = create_engine(f"sqlite:///{tmp_path / 'db.sqlite'}", future=True)
    Base.metadata.create_all(engine)
    Session = sessionmaker(bind=engine, future=True, expire_on_commit=False)
    with Session() as session:
        session.add(Project(id=1, name="Demo", code="DEMO"))
        session.add(
            Employee(id=1, full_name="Адриан Сергей Александрович", is_available_in_bitrix=True)
        )
        session.commit()
        yield session


def build_memo_docx(path: Path) -> Path:
    doc = Document()
    for line in FIXTURE_TEXT.read_text(encoding="utf-8").splitlines():
        doc.add_paragraph(line)
    doc.save(path)
    return path


def _result(tmp_path: Path):
    doc = parse_docx(str(build_memo_docx(tmp_path / "memo_m026_26.docx")))
    parser, choice = ParserRegistry().choose(doc)
    return doc, parser, choice, parser.parse(doc)


def test_memo_parser_real_docx_regression(db, tmp_path):
    doc, parser, choice, result = _result(tmp_path)
    assert parser.parser_type == "memo_protocol"
    assert choice.confidence >= 0.9
    assert len(result.tasks) == 9
    assert [t.task_number for t in result.tasks] == [str(i) for i in range(1, 10)]
    all_text = "\n".join(t.title for t in result.tasks)
    assert "Рабочей встречи по теме" not in all_text
    assert "создание Единого центра компетенций" not in all_text
    assert "г. Санкт-Петербург" not in all_text
    assert "ОТМЕТИЛИ" not in all_text
    assert "Мемо подготовил" not in result.tasks[-1].source_text
    assert (
        result.tasks[0].title
        == "Сопоставить итоги внутреннего рейтинга по качеству предоставляемых услуг (перевозка, питание, проживание) для работников ООО «ГСП-Технологии» на объектах присутствия с данными по тепловой карте ВЖГ. Разработать план корректирующих мероприятий."
    )
    assert result.tasks[0].assignee_raw == "Адриан С.А., Грибачев С.П."
    assert result.tasks[1].assignee_raw == "Прокофьев Д.Ю., Адриан С.А."
    assert result.tasks[6].assignee_raw == "Адриан С.А., Прокофьев Д.Ю."
    assert result.tasks[7].assignee_raw == "Грибачев С.П., Прокофьев Д.Ю."
    assert all(t.deadline for t in result.tasks)
    assert [t.deadline for t in result.tasks[:4]] == [
        "2026-05-12",
        "2026-05-15",
        "2026-05-12",
        "2026-05-14",
    ]
    assert result.tasks[2].block_raw == "Организация работ на проекте Усть-Луга"
    assert {t.block_raw for t in result.tasks[4:7]} == {"Укомплектование ЛР"}
    assert {t.block_raw for t in result.tasks[7:9]} == {"Организация работ ЦТЗ"}
    assert not any("Срок не распознан" in w for t in result.tasks for w in t.warnings)
    assert all(t.parsing_confidence >= 0.9 for t in result.tasks)
    assert [t.title for t in MemoProtocolParser().parse(doc).tasks] == [
        t.title for t in result.tasks
    ]


def test_memo_deadline_label_variants():
    doc = ParsedDocument(
        "x.docx",
        elements=[
            ParsedElement(0, "ИТОГИ", "paragraph", SourceLocation()),
            ParsedElement(1, "ОТМЕТИЛИ", "paragraph", SourceLocation()),
            ParsedElement(2, "РЕШИЛИ:", "paragraph", SourceLocation()),
            ParsedElement(3, "1. Сделать первое поручение", "paragraph", SourceLocation()),
            ParsedElement(4, "Исполнитель: Иванов И.И.", "paragraph", SourceLocation()),
            ParsedElement(5, "Срок:", "paragraph", SourceLocation()),
            ParsedElement(6, "12.05.2026", "paragraph", SourceLocation()),
            ParsedElement(7, "2) Сделать второе поручение", "paragraph", SourceLocation()),
            ParsedElement(8, "Ответственный: Петров П.П.", "paragraph", SourceLocation()),
            ParsedElement(9, "Срок 12.05.26", "paragraph", SourceLocation()),
        ],
    )
    result = MemoProtocolParser().parse(doc)
    assert [t.deadline for t in result.tasks] == ["2026-05-12", "2026-05-12"]


def test_universal_parser_still_parses_basic_fixture():
    parsed = UniversalProtocolParser().parse(
        ParsedDocument(
            "demo.docx",
            [],
            [
                ParsedElement(0, "Protocol", "paragraph", SourceLocation()),
                ParsedElement(1, "1. Do work до 12.05.2026", "paragraph", SourceLocation()),
            ],
        )
    )
    assert parsed.tasks[0].task_number == "1"


def test_import_preview_uses_memo_protocol(db, tmp_path):
    session = create_preview_session(
        db, 1, upload(build_memo_docx(tmp_path / "memo_m026_26.docx"), "memo_m026_26.docx")
    )
    assert session.parser_type == "memo_protocol"
    assert session.parsed_payload["parser_choice"]["parser_type"] == "memo_protocol"
    assert len(session.parsed_payload["tasks"]) == 9
    assert not any("Срок не распознан" in w for w in session.warnings_payload)


def test_empty_memo_cannot_be_confirmed(db, tmp_path):
    doc = Document()
    for line in [
        "МЕМО",
        "РЕШИЛИ:",
        "Исполнитель: Иванов И.И.",
        "Срок: 12.05.2026",
        "Мемо подготовил: Петров П.П.",
    ]:
        doc.add_paragraph(line)
    path = tmp_path / "empty_memo.docx"
    doc.save(path)

    session = create_preview_session(db, 1, upload(path, path.name))

    assert session.parsed_payload["tasks"] == []
    assert session.errors_payload == [
        "Не удалось распознать ни одного поручения МЕМО. "
        "Проверьте раздел «РЕШИЛИ» и формат «текст поручения → исполнитель → срок»."
    ]
    with pytest.raises(HTTPException, match="не распознано ни одного поручения"):
        confirm_session(db, session)


def build_structured_memo_docx(path: Path) -> Path:
    doc = Document()
    for line in [
        "МЕМО М-026/26",
        "03.05.2026, г. Санкт-Петербург",
        "ОТМЕТИЛИ",
        "создание Единого центра компетенций",
        "Р Е Ш И Л И :",
        "1. Уточнить потребность 04.05.26 по результатам совещания",
        "Исполнитель: Адриан С.А., Грибачев С.П.",
        "Срок: 12.05.2026",
        "2. Подготовить план мероприятий",
        "Исполнитель: Прокофьев Д.Ю., Адриан С.А.",
        "Срок 15.05.2026",
        "Блок 1: Организация работ на проекте Усть-Луга",
        "3. Согласовать порядок взаимодействия",
        "Исполнитель: Иванов И.И.",
        "Срок – 12.05.2026",
        "4. Проверить доступы",
        "Исполнитель: Петров П.П.",
        "Срок до 14.05.2026",
        "Блок 2: Укомплектование ЛР",
        "5. Сформировать список потребностей",
        "Исполнитель: Сидоров С.С.",
        "Срок: 16.05.2026",
        "6. Провести сверку состава",
        "Исполнитель: Смирнов С.С.",
        "Срок: 17.05.2026",
        "7. Подготовить отчет",
        "Исполнитель: Адриан С.А., Прокофьев Д.Ю.",
        "Срок: 18.05.2026",
        "Блок 3: Организация работ ЦТЗ",
        "8. Назначить ответственных",
        "Исполнитель: Грибачев С.П., Прокофьев Д.Ю.",
        "Срок: 19.05.2026",
        "9. Направить итоговое письмо",
        "Исполнитель: Орлов О.О.",
        "Срок: 20.05.2026",
        "Мемо подготовил",
        "Автор Документа",
    ]:
        doc.add_paragraph(line)
    doc.save(path)
    return path


def test_import_preview_endpoint_uses_memo_protocol_for_real_docx_path(db, tmp_path):
    from fastapi.testclient import TestClient

    from app.db.session import get_db
    from app.main import app

    def override_get_db():
        yield db

    app.dependency_overrides[get_db] = override_get_db
    try:
        path = build_structured_memo_docx(tmp_path / "02_Мемо_М-026_26.docx")
        with path.open("rb") as f:
            response = TestClient(app, follow_redirects=False).post(
                "/protocols/import/preview",
                data={"project_id": "1", "parser_type": "universal"},
                files={
                    "file": (
                        "02_Мемо_М-026_26 от 03.05.2026.docx",
                        f,
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        assert response.status_code == 303
        session_id = int(response.headers["location"].split("/")[3])
        session = db.get(ImportSession, session_id)
    finally:
        app.dependency_overrides.clear()
    payload = session.parsed_payload
    assert session.parser_type == "memo_protocol"
    assert payload["parser_choice"]["parser_type"] == "memo_protocol"
    diagnostics = payload["metadata"]["diagnostics"]
    assert diagnostics["element_count"] >= 30
    assert len(diagnostics["first_30_normalized_elements"]) == 30
    assert diagnostics["decision_section_index"] == 4
    assert diagnostics["memo_prepared_index"] == 35
    assert diagnostics["elements_between_decision_and_prepared"] == 30
    assert diagnostics["rejection_reasons"] == []
    assert len(payload["tasks"]) == 9
    titles = "\n".join(t["title"] for t in payload["tasks"])
    assert "МЕМО" not in titles
    assert "Исполнитель:" not in titles
    assert "Срок" not in titles
    assert "Блок" not in titles
    assert "Мемо подготовил" not in titles
    assert "Автор Документа" not in titles
    assert payload["tasks"][0]["deadline"] == "2026-05-12"
    assert payload["tasks"][0]["assignee_raw"] == "Адриан С.А., Грибачев С.П."
    assert payload["tasks"][0]["title"].endswith("по результатам совещания")
    assert payload["tasks"][0]["deadline"] != "2026-05-04"
    assert payload["tasks"][2]["block_raw"] == "Организация работ на проекте Усть-Луга"
    assert payload["tasks"][4]["block_raw"] == "Укомплектование ЛР"
    assert payload["tasks"][7]["block_raw"] == "Организация работ ЦТЗ"
    assert not any("Специализированный парсер МЕМО не применён" in w for w in session.warnings_payload)


def test_real_tabular_memo_binary_fixture_regression(db, tmp_path):
    encoded_fixture = Path("tests/fixtures/real_tabular_memo.docx.b64")
    path = tmp_path / "real_tabular_memo.docx"
    path.write_bytes(base64.b64decode(encoded_fixture.read_bytes()))
    document = parse_docx(str(path))
    result = MemoProtocolParser().parse(document)

    assert len(result.tasks) == 9
    assert [task.task_number for task in result.tasks] == [str(i) for i in range(1, 10)]
    assert len(result.sections) == 4
    assert [section.title for section in result.sections] == [
        "Сварочный трест",
        "Монтажный участок",
        "Проектный офис",
        "Пусконаладочные работы",
    ]
    assert {task.block_raw for task in result.tasks} == {section.title for section in result.sections}
    assert all(task.assignee_raw for task in result.tasks)
    assert all(task.deadline for task in result.tasks)
    assert result.tasks[0].assignee_raw == "Иванов И.И."
    assert result.tasks[3].assignee_raw == "Смирнов С.С."
    assert result.tasks[5].assignee_raw == "Волков В.В."
    assert result.tasks[-1].deadline == "2026-08-20"
    titles = "\n".join(task.title for task in result.tasks)
    assert "#Сварочный трест" not in titles
    assert "Исполнитель" not in titles
    assert "Срок" not in titles
    assert "Мемо подготовил" not in titles
    assert result.errors == []


def test_confirmation_is_blocked_when_document_assignee_was_not_recognized(db, tmp_path):
    doc = Document()
    doc.add_paragraph("МЕМО")
    doc.add_paragraph("РЕШИЛИ:")
    doc.add_paragraph("1. Выполнить обязательное поручение")
    doc.add_paragraph("Исполнитель:")
    doc.add_paragraph("Срок: 12.08.2026")
    path = tmp_path / "missing-assignee.docx"
    doc.save(path)

    session = create_preview_session(db, 1, upload(path, path.name))

    assert session.parsed_payload["tasks"][0]["assignee_raw"] is None
    assert "исполнители указаны в документе" in session.errors_payload[0]
    with pytest.raises(HTTPException, match="ошибки распознавания исполнителей"):
        confirm_session(db, session)
